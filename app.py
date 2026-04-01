import io
import os
import re
import time
import uuid
import zipfile
import threading
import traceback

import cv2
import faiss
import numpy as np
import pandas as pd
import pdfplumber
import docx as python_docx
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageFile
from pdf2image import convert_from_path
from sentence_transformers import SentenceTransformer
from openai import OpenAI
from flask import Flask, request, render_template, jsonify
import easyocr


UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
ImageFile.LOAD_TRUNCATED_IMAGES = True

POPPLER_PATH = os.getenv(
    "POPPLER_PATH",
    r"C:\Users\admin\Downloads\Release-25.12.0-0\poppler-25.12.0\Library\bin"
)
if not os.path.isdir(POPPLER_PATH):
    raise RuntimeError(
        f"Poppler not found at: {POPPLER_PATH}\n"
        "Set POPPLER_PATH environment variable to your Poppler bin folder."
    )

LM_BASE_URL          = os.getenv("LM_BASE_URL",          "http://127.0.0.1:1234/v1")
LM_BASE_URL_FALLBACK = os.getenv("LM_BASE_URL_FALLBACK", "http://127.0.0.1:1234/api/v1")
LLM_MODEL            = os.getenv("LLM_MODEL",            "meta-llama-3-8b-instruct")
EMBED_MODEL          = os.getenv("EMBED_MODEL",          "sentence-transformers/all-MiniLM-L6-v2")

OCR_DPI              = int(os.getenv("OCR_DPI",              "200"))
TABLE_DPI            = int(os.getenv("TABLE_DPI",            "300"))
POPPLER_THREADS      = int(os.getenv("POPPLER_THREADS",      "4"))
SCANNED_PAGE_CHARS   = int(os.getenv("SCANNED_PAGE_CHARS",   "80"))
SCANNED_DOC_RATIO    = float(os.getenv("SCANNED_DOC_RATIO",  "0.4"))
SCANNED_SAMPLE_PAGES = int(os.getenv("SCANNED_SAMPLE_PAGES", "8"))
RETRIEVAL_TOP_K      = int(os.getenv("RETRIEVAL_TOP_K",      "6"))
RETRIEVAL_CHAR_LIMIT = int(os.getenv("RETRIEVAL_CHAR_LIMIT", "4000"))
RETRIEVAL_MAX_DIST   = float(os.getenv("RETRIEVAL_MAX_DIST", "0.8"))
MAX_UPLOAD_MB        = int(os.getenv("MAX_UPLOAD_MB",        "500"))
EXCEL_PATH           = os.getenv("EXCEL_PATH", os.path.join(UPLOAD_FOLDER, "extracted_tables.xlsx"))



app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_MB * 1024 * 1024



ocr_reader = easyocr.Reader(["en"], gpu=False, verbose=False)

llm          = OpenAI(base_url=LM_BASE_URL,          api_key="lm-studio")
llm_fallback = OpenAI(base_url=LM_BASE_URL_FALLBACK, api_key="lm-studio")

embedder  = SentenceTransformer(EMBED_MODEL, device="cpu")
EMBED_DIM = embedder.get_sentence_embedding_dimension()



INDEX        = faiss.IndexFlatL2(EMBED_DIM)
CHUNKS       = []
INDEX_LOCK   = threading.Lock()
FILE_META    = {}
FILE_CHUNKS  = {}
TABLE_CACHE  = {}
RECENT_FILES = []
JOBS         = {}
JOBS_LOCK    = threading.Lock()
MAX_JOB_S    = int(os.getenv("MAX_JOB_SECONDS", "900"))



def deskew(gray):
    try:
        edges = cv2.Canny(gray, 50, 150, apertureSize=3)
        lines = cv2.HoughLinesP(
            edges, 1, np.pi / 180, threshold=100,
            minLineLength=gray.shape[1] // 4, maxLineGap=20
        )
        if lines is None:
            return gray
        angles = []
        for x1, y1, x2, y2 in lines[:, 0]:
            angle = np.degrees(np.arctan2(y2 - y1, x2 - x1))
            if abs(angle) < 5:
                angles.append(angle)
        if not angles:
            return gray
        angle = float(np.median(angles))
        if abs(angle) < 0.3:
            return gray
        h, w = gray.shape
        M = cv2.getRotationMatrix2D((w / 2, h / 2), angle, 1.0)
        return cv2.warpAffine(gray, M, (w, h),
                              flags=cv2.INTER_CUBIC,
                              borderMode=cv2.BORDER_REPLICATE)
    except Exception:
        return gray


def preprocess_for_ocr(img_bgr):
    try:
        gray  = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        gray  = clahe.apply(gray)
        gray  = cv2.fastNlMeansDenoising(gray, h=10,
                    templateWindowSize=7, searchWindowSize=21)
        gray  = deskew(gray)
        gray  = cv2.resize(gray, None, fx=2.0, fy=2.0,
                    interpolation=cv2.INTER_CUBIC)
        blur  = cv2.GaussianBlur(gray, (0, 0), 3)
        gray  = cv2.addWeighted(gray, 1.5, blur, -0.5, 0)
        bw    = cv2.adaptiveThreshold(gray, 255,
                    cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                    cv2.THRESH_BINARY, 21, 8)
        return bw
    except Exception:
        return cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)


def preprocess_for_table(img_bgr):
    try:
        gray  = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        gray  = clahe.apply(gray)
        gray  = cv2.fastNlMeansDenoising(gray, h=8,
                    templateWindowSize=7, searchWindowSize=21)
        gray  = deskew(gray)
        gray  = cv2.resize(gray, None, fx=2.0, fy=2.0,
                    interpolation=cv2.INTER_CUBIC)
        bw    = cv2.adaptiveThreshold(gray, 255,
                    cv2.ADAPTIVE_THRESH_MEAN_C,
                    cv2.THRESH_BINARY, 31, 10)
        hk    = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
        vk    = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
        hl    = cv2.morphologyEx(bw, cv2.MORPH_CLOSE, hk)
        vl    = cv2.morphologyEx(bw, cv2.MORPH_CLOSE, vk)
        borders = cv2.bitwise_and(hl, vl)
        result  = cv2.bitwise_or(
            cv2.bitwise_and(bw, cv2.bitwise_not(borders)), borders
        )
        return result
    except Exception:
        return cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)




def ocr_image(img_gray):
    try:
        results = ocr_reader.readtext(img_gray, detail=0, paragraph=True)
        return " ".join(str(r) for r in results).strip()
    except Exception as e:
        print(f"[EasyOCR] {e}")
        return ""




def get_page_count(path):
    try:
        with pdfplumber.open(path) as pdf:
            return len(pdf.pages)
    except Exception:
        return 0


def page_has_text(page):
    return len((page.extract_text() or "").strip()) >= SCANNED_PAGE_CHARS


def pil_to_bgr(pil_img):
    arr = np.array(pil_img)
    if len(arr.shape) == 2:
        return cv2.cvtColor(arr, cv2.COLOR_GRAY2BGR)
    return cv2.cvtColor(arr, cv2.COLOR_RGB2BGR)


def convert_pages(path, first, last, dpi):
    return convert_from_path(
        path, dpi=dpi,
        poppler_path=POPPLER_PATH,
        first_page=first, last_page=last,
        thread_count=max(1, POPPLER_THREADS)
    )



def table_list_to_md(tbl, page_no=0):
    if not tbl:
        return ""
    clean = [[(c or "").replace("\n", " ") for c in row] for row in tbl]
    n     = max(len(r) for r in clean)
    clean = [r + [""] * (n - len(r)) for r in clean]
    label = f"[Table | Page {page_no}]\n" if page_no else ""
    md    = label
    md   += "| " + " | ".join(clean[0]) + " |\n"
    md   += "| " + " | ".join(["---"] * n) + " |\n"
    for r in clean[1:]:
        md += "| " + " | ".join(r) + " |\n"
    return md.strip()


def quality_ok(md):
    if not md or "|" not in md:
        return False
    data_rows = [l for l in md.splitlines()
                 if "|" in l and not re.fullmatch(r"[\s|:\-]+", l)]
    return len(data_rows) >= 3




def extract_pdf_text(path, max_pages=0):
    full_text = ""
    try:
        with pdfplumber.open(path) as pdf:
            pages = pdf.pages if max_pages == 0 else pdf.pages[:max_pages]

            scanned_nos = [i + 1 for i, p in enumerate(pages)
                           if not page_has_text(p)]

            ocr_images = {}
            if scanned_nos:
                try:
                    imgs    = convert_pages(path, scanned_nos[0],
                                            scanned_nos[-1], OCR_DPI)
                    img_idx = 0
                    for pg in range(scanned_nos[0], scanned_nos[-1] + 1):
                        if img_idx < len(imgs):
                            ocr_images[pg] = imgs[img_idx]
                            img_idx += 1
                except Exception as e:
                    print(f"[PDF batch convert] {e}")

            for i, page in enumerate(pages, start=1):
                full_text += f"\n--- Page {i} ---\n"
                if page_has_text(page):
                    full_text += (page.extract_text() or "")
                    for tbl in (page.extract_tables() or []):
                        md = table_list_to_md(tbl, i)
                        if md:
                            full_text += f"\n{md}\n"
                else:
                    pil = ocr_images.get(i)
                    if pil is not None:
                        img  = pil_to_bgr(pil)
                        gray = preprocess_for_ocr(img)
                        full_text += ocr_image(gray)

    except Exception as e:
        print(f"[extract_pdf_text] {e}")
    return full_text.strip()



def extract_tables_text_layer(path, start=1, end=0):
    result = ""
    try:
        with pdfplumber.open(path) as pdf:
            total = len(pdf.pages)
            s     = max(1, start)
            e     = total if end == 0 else min(end, total)
            for i, page in enumerate(pdf.pages[s - 1: e], start=s):
                for tbl in (page.extract_tables() or []):
                    md = table_list_to_md(tbl, i)
                    if md:
                        result += md + "\n\n"
    except Exception as e:
        print(f"[text-layer tables] {e}")
    return result.strip()


def extract_tables_ocr(path, start=1, end=0):
    result = ""
    try:
        with pdfplumber.open(path) as pdf:
            total = len(pdf.pages)
        s    = max(1, start)
        e    = total if end == 0 else min(end, total)
        imgs = convert_pages(path, s, e, TABLE_DPI)
        for rel, pil in enumerate(imgs):
            page_no = s + rel
            try:
                img  = pil_to_bgr(pil)
                gray = preprocess_for_table(img)
                text = ocr_image(gray)
                if text.strip():
                    md = format_table_with_llm(text, page_no)
                    if md and "|" in md:
                        result += f"[Table | Page {page_no}]\n{md}\n\n"
            except Exception as e:
                print(f"[OCR table page {page_no}] {e}")
    except Exception as e:
        print(f"[extract_tables_ocr] {e}")
    return result.strip()


def find_first_table_page(path, limit=60):
    try:
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages[:limit]):
                if page.extract_tables():
                    return i + 1
    except Exception:
        pass
    return None


def extract_tables(path, start=1, end=0):
    """
    Always scan ALL pages in range.
    Text-layer first on every page; OCR fallback only on scanned pages.
    """
    text_md = extract_tables_text_layer(path, start, end)

    # Find scanned pages (no text layer) and run OCR only on those
    ocr_md = ""
    try:
        with pdfplumber.open(path) as pdf:
            total = len(pdf.pages)
            s = max(1, start)
            e = total if end == 0 else min(end, total)
            scanned_pages = [
                i for i in range(s, e + 1)
                if not page_has_text(pdf.pages[i - 1])
            ]
        for pg in scanned_pages:
            pg_md = extract_tables_ocr(path, start=pg, end=pg)
            if pg_md:
                ocr_md += pg_md + "\n\n"
    except Exception as e:
        print(f"[extract_tables ocr pass] {e}")

    combined = (text_md + "\n\n" + ocr_md).strip()
    return combined if combined else ""


def extract_docx_text(path):
    text = ""
    try:
        doc = python_docx.Document(path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for tbl in doc.tables:
            rows = [[c.text.replace("\n", " ") for c in row.cells]
                    for row in tbl.rows]
            if rows:
                n    = max(len(r) for r in rows)
                rows = [r + [""] * (n - len(r)) for r in rows]
                text += "\n[TABLE DATA]\n"
                text += "| " + " | ".join(rows[0]) + " |\n"
                text += "| " + " | ".join(["---"] * n) + " |\n"
                for r in rows[1:]:
                    text += "| " + " | ".join(r) + " |\n"
        with zipfile.ZipFile(path, "r") as z:
            for name in z.namelist():
                if name.startswith("word/media/"):
                    try:
                        data = z.read(name)
                        pil  = Image.open(io.BytesIO(data)).convert("RGB")
                        img  = pil_to_bgr(pil)
                        gray = preprocess_for_ocr(img)
                        text += "\n" + ocr_image(gray) + "\n"
                    except Exception as e:
                        print(f"[DOCX image OCR] {e}")
    except Exception as e:
        print(f"[extract_docx_text] {e}")
    return text.strip()


def extract_image_text(path):
    try:
        pil  = Image.open(path).convert("RGB")
        img  = pil_to_bgr(pil)
        gray = preprocess_for_ocr(img)
        return ocr_image(gray)
    except Exception as e:
        print(f"[extract_image_text] {e}")
        return ""



def extract_text(path, max_pages=0):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_pdf_text(path, max_pages)
    if ext == ".docx":
        return extract_docx_text(path)
    if ext in (".png", ".jpg", ".jpeg"):
        return extract_image_text(path)
    return ""


def query_llm(prompt, max_tokens=512):
    for base_url, client in [(LM_BASE_URL, llm),
                              (LM_BASE_URL_FALLBACK, llm_fallback)]:
        try:
            r = client.chat.completions.create(
                model=LLM_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0,
                max_tokens=max_tokens,
                timeout=120,
                extra_body={"repetition_penalty": 1.1,
                            "repeat_penalty":     1.1},
            )
            return r.choices[0].message.content.strip()
        except Exception as e:
            print(f"[LLM {base_url}] {e}")
    return (
        "Could not reach LM Studio. "
        "Open LM Studio → Developer → Local Server → set Status to Running."
    )


def format_table_with_llm(ocr_text, page_no):
    prompt = (
        "You are a data extraction engine. Output ONLY a Markdown table.\n"
        "Rules:\n"
        "1. Copy values EXACTLY as they appear — do NOT paraphrase or infer.\n"
        "2. Keep column count consistent per table.\n"
        "3. Empty cells stay blank (just ||).\n"
        "4. Include a header row and a separator row (---).\n"
        "5. If no table structure exists, reply exactly: NO_TABLE\n\n"
        f"TEXT FROM PAGE {page_no}:\n{ocr_text}"
    )
    result = query_llm(prompt, max_tokens=512)
    if "NO_TABLE" in result:
        return ""
    return result


def chunk_text(text, size=1000, overlap=200):
    chunks, current = [], ""
    for para in re.split(r"\n{2,}", text):
        para = para.strip()
        if not para:
            continue
        if len(current) + len(para) + 2 <= size:
            current = (current + "\n\n" + para).strip()
        else:
            if current:
                chunks.append(current)
            if len(para) > size:
                for i in range(0, len(para), size - overlap):
                    chunks.append(para[i: i + size])
                current = ""
            else:
                current = para
    if current:
        chunks.append(current)
    return chunks


def index_text(text, file_id):
    if not text.strip():
        return 0
    chunks = chunk_text(text)
    embs   = embedder.encode(chunks, show_progress_bar=False,
                             normalize_embeddings=True)
    with INDEX_LOCK:
        INDEX.add(np.array(embs).astype("float32"))
        CHUNKS.extend(chunks)
        FILE_CHUNKS.setdefault(file_id, []).extend(chunks)
    return len(chunks)


def reset_index(file_ids):
    global INDEX, CHUNKS, FILE_CHUNKS
    with INDEX_LOCK:
        INDEX       = faiss.IndexFlatL2(EMBED_DIM)
        CHUNKS      = []
        FILE_CHUNKS = {fid: [] for fid in file_ids}


def retrieve(query, top_k=RETRIEVAL_TOP_K):
    with INDEX_LOCK:
        if INDEX.ntotal == 0:
            return []
        q    = embedder.encode([query], normalize_embeddings=True).astype("float32")
        D, I = INDEX.search(q, top_k)
    return [(CHUNKS[i], float(d)) for d, i in zip(D[0], I[0]) if i < len(CHUNKS)]


def build_context(pairs):
    parts, total = [], 0
    for chunk, dist in pairs:
        if dist > RETRIEVAL_MAX_DIST:
            continue
        if total + len(chunk) > RETRIEVAL_CHAR_LIMIT and parts:
            break
        parts.append(chunk)
        total += len(chunk)
    return "\n\n".join(parts)



def md_to_df(md):
    lines = [l for l in md.splitlines() if l.strip()]
    if not lines:
        return None
    if lines[0].startswith("[Table"):
        lines = lines[1:]
    lines = [l for l in lines if not re.fullmatch(r"[\s|:\-]+", l)]
    if len(lines) < 2:
        return None
    rows = [[c.strip() for c in l.strip().strip("|").split("|")] for l in lines]
    n    = max(len(r) for r in rows)
    rows = [r + [""] * (n - len(r)) for r in rows]
    df   = pd.DataFrame(rows[1:], columns=rows[0])
    df   = df.dropna(how="all").reset_index(drop=True)
    df.columns = [str(c).strip() for c in df.columns]
    return df if not df.empty else None


def save_excel(md_text, path=EXCEL_PATH):
    if not md_text or "|" not in md_text:
        return None
    blocks = [b.strip() for b in re.split(r"\n\s*\n", md_text) if b.strip()]
    dfs    = [md_to_df(b) for b in blocks if "|" in b]
    dfs    = [d for d in dfs if d is not None]
    if not dfs:
        return None
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    HDR_FILL  = PatternFill("solid", fgColor="1F4E79")
    HDR_FONT  = Font(bold=True, color="FFFFFF", size=10)
    HDR_ALIGN = Alignment(horizontal="center", vertical="center", wrap_text=True)
    DAT_ALIGN = Alignment(horizontal="left",   vertical="top",    wrap_text=True)
    THIN      = Side(style="thin", color="D0D7E3")
    BORDER    = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
    for idx, df in enumerate(dfs, 1):
        ws = wb.create_sheet(f"Table{idx}")
        for ci, col in enumerate(df.columns, 1):
            c = ws.cell(1, ci, str(col))
            c.fill = HDR_FILL; c.font = HDR_FONT
            c.alignment = HDR_ALIGN; c.border = BORDER
        for ri, row in df.iterrows():
            for ci, val in enumerate(row, 1):
                c = ws.cell(ri + 2, ci, str(val) if val is not None else "")
                c.alignment = DAT_ALIGN; c.border = BORDER
        for col in ws.columns:
            w = max((len(str(c.value or "")) for c in col), default=8)
            ws.column_dimensions[get_column_letter(col[0].column)].width = min(w + 4, 60)
        ws.freeze_panes = "A2"
        ws.row_dimensions[1].height = 22
    wb.save(path)
    return path


def parse_page_hint(text):
    m = re.search(r"page\s*(\d+)", text, re.IGNORECASE)
    return max(1, int(m.group(1))) if m else None


def parse_page_range(text):
    m = re.search(r"pages?\s*(\d+)\s*[-to]+\s*(\d+)", text, re.IGNORECASE)
    if m:
        s, e = max(1, int(m.group(1))), int(m.group(2))
        return s, max(s, e)
    return None


def get_tables_cached(file_id, path, start=1, end=0):
    key = f"{start}_{end}"
    if file_id in TABLE_CACHE and key in TABLE_CACHE[file_id]:
        return TABLE_CACHE[file_id][key]
    md = extract_tables(path, start, end)
    TABLE_CACHE.setdefault(file_id, {})[key] = md
    return md


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload():
    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "No files uploaded"})

    job_id = str(uuid.uuid4())
    saved  = []

    for f in files:
        path = os.path.join(UPLOAD_FOLDER, f"{uuid.uuid4()}_{f.filename}")
        f.save(path)
        fid  = str(uuid.uuid4())
        ext  = os.path.splitext(path)[1].lower()
        saved.append((path, fid, f.filename, ext))
        RECENT_FILES.append({"id": fid, "path": path,
                              "name": f.filename, "ts": time.time()})
        FILE_META[fid] = {
            "path": path, "ext": ext,
            "page_count": get_page_count(path) if ext == ".pdf" else 0,
        }

    if len(RECENT_FILES) > 20:
        del RECENT_FILES[:-20]

    reset_index([fid for _, fid, _, _ in saved])

    with JOBS_LOCK:
        JOBS[job_id] = {"status": "queued",
                        "message": f"Queued {len(saved)} file(s)",
                        "start_ts": time.time()}

    def worker():
        with JOBS_LOCK:
            JOBS[job_id]["status"] = "running"
        try:
            for i, (path, fid, fname, ext) in enumerate(saved, 1):
                with JOBS_LOCK:
                    JOBS[job_id]["message"] = f"Indexing {i}/{len(saved)}: {fname}"
                text = extract_text(path) or ""
                if text:
                    index_text(text, fid)
            with JOBS_LOCK:
                JOBS[job_id]["status"]  = "done"
                JOBS[job_id]["message"] = f"{len(saved)} document(s) ready"
        except Exception as e:
            print(traceback.format_exc())
            with JOBS_LOCK:
                JOBS[job_id]["status"]  = "error"
                JOBS[job_id]["message"] = f"Failed: {e}"

    threading.Thread(target=worker, daemon=True).start()
    return jsonify({"job_id": job_id, "status": "queued",
                    "message": f"Queued {len(saved)} file(s)"})


@app.route("/status/<job_id>")
def status(job_id):
    with JOBS_LOCK:
        job = JOBS.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    r = dict(job)
    if r.get("status") == "running" and r.get("start_ts"):
        elapsed = int(time.time() - r["start_ts"])
        r["elapsed_seconds"] = elapsed
        if elapsed > MAX_JOB_S:
            r["status"]  = "error"
            r["message"] = f"Timed out after {MAX_JOB_S}s"
    return jsonify(r)


@app.route("/chat", methods=["POST"])
def chat():
    question = (request.json or {}).get("message", "").strip()
    if not question:
        return jsonify({"answer": "Please provide a question."})

    q_lower = question.lower()

    # Table path
    if any(w in q_lower for w in ("table", "extract", "tabular")):
        if RECENT_FILES:
            latest = RECENT_FILES[-1]
            if latest["path"].lower().endswith(".pdf"):
                pr = parse_page_range(question)
                ph = None if pr else parse_page_hint(question)
                if pr:
                    start, end = pr
                elif ph:
                    start, end = ph, ph
                else:
                    start, end = 1, 0
                md = get_tables_cached(latest["id"], latest["path"], start, end)
                if md:
                    save_excel(md)
                    return jsonify({"answer": md})
                return jsonify({"answer": "No tables detected in the document."})

    pairs   = retrieve(question)
    context = build_context(pairs)

    if not context and RECENT_FILES:
        latest = RECENT_FILES[-1]
        if not FILE_CHUNKS.get(latest["id"]):
            text = extract_text(latest["path"]) or ""
            if text:
                index_text(text, latest["id"])
                pairs   = retrieve(question)
                context = build_context(pairs)

    if not context:
        return jsonify({"answer": "Not mentioned in the document."})

    prompt = (
        "You are a strict document assistant. "
        "Use ONLY the information explicitly stated in the CONTEXT. "
        "Do NOT use prior knowledge, infer, or guess. "
        "If the answer needs a table: output ONLY a Markdown table with header and separator rows. "
        "If plain text: output ONLY plain text. "
        "If not in CONTEXT: reply exactly — Not mentioned in the document.\n\n"
        f"CONTEXT:\n{context}\n\n"
        f"QUESTION:\n{question}"
    )
    return jsonify({"answer": query_llm(prompt)})


@app.route("/tables", methods=["POST"])
def tables():
    if not RECENT_FILES:
        return jsonify({"answer": "No document uploaded yet."})
    latest = RECENT_FILES[-1]
    if not latest["path"].lower().endswith(".pdf"):
        return jsonify({"answer": "Latest file is not a PDF."})

    payload = request.get_json(silent=True) or {}
    query   = (payload.get("message") or "").strip()
    pr      = parse_page_range(query)
    ph      = None if pr else parse_page_hint(query)

    if pr:
        start, end = pr
    elif ph:
        start, end = ph, ph
    else:
        start, end = 1, 0

    try:
        md = get_tables_cached(latest["id"], latest["path"], start, end)
        if md:
            save_excel(md)
            return jsonify({"answer": md})
        return jsonify({"answer": "No tables detected in the document."})
    except Exception as e:
        return jsonify({"answer": f"Table extraction failed: {e}"})


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=False, use_reloader=False)
